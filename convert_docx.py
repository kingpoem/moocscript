"""Script to convert Markdown files to DOCX format."""

import argparse
import re
import tempfile
import shutil
from pathlib import Path
from typing import Optional
from urllib.parse import urlparse

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    import requests
except ImportError:
    print("Error: Required packages not installed.")
    print("Please install: uv sync")
    print("Or manually: pip install python-docx requests")
    exit(1)


def find_markdown_files(markdown_dir: Path) -> dict:
    """Find all Markdown files organized by course.
    
    Args:
        markdown_dir: Directory containing course subdirectories with Markdown files
        
    Returns:
        Dictionary mapping course names to lists of markdown files
    """
    courses_data = {}
    
    if not markdown_dir.exists():
        print(f"Markdown directory not found: {markdown_dir}")
        print("  Please run convert_markdown.py first to generate Markdown files")
        return courses_data
    
    # Iterate through course directories
    for course_dir in markdown_dir.iterdir():
        if not course_dir.is_dir():
            continue
        
        course_name = course_dir.name
        markdown_files = []
        
        # Find all Markdown files
        for md_file in course_dir.glob("*.md"):
            markdown_files.append(md_file)
        
        if markdown_files:
            courses_data[course_name] = markdown_files
    
    return courses_data


def download_image(image_url: str, cache_dir: Optional[Path] = None) -> Optional[Path]:
    """Download an image from URL and save it to cache directory.
    
    Args:
        image_url: URL of the image to download
        cache_dir: Directory to cache downloaded images (optional)
        
    Returns:
        Path to the downloaded image file, or None if download failed
    """
    if not image_url:
        return None
    
    try:
        # Create cache directory if provided
        if cache_dir:
            cache_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate filename from URL
            parsed_url = urlparse(image_url)
            filename = Path(parsed_url.path).name
            if not filename or '.' not in filename:
                # Generate extension from content type or use default
                filename = f"image_{hash(image_url) % 100000}.jpg"
            cache_file = cache_dir / filename
            
            # Check if already cached
            if cache_file.exists():
                return cache_file
            
            # Download image
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(image_url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Save to cache
            with open(cache_file, 'wb') as f:
                f.write(response.content)
            
            return cache_file
        else:
            # Use temporary file if no cache directory
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(image_url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Create temporary file
            suffix = Path(urlparse(image_url).path).suffix or '.jpg'
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            temp_file.write(response.content)
            temp_file.close()
            
            return Path(temp_file.name)
    except Exception as e:
        print(f"  Warning: Failed to download image {image_url}: {str(e)}")
        return None


def add_text_with_images(doc: Document, content: str, image_cache_dir: Optional[Path] = None, 
                         font_size: Pt = None, font_name: str = '宋体', bold: bool = False,
                         paragraph: Optional = None, inline_images: bool = False,
                         paragraph_indent: Optional[Inches] = None) -> list:
    """Add text content to a paragraph, processing any images inline.
    
    Args:
        doc: Document object
        content: Text content that may contain image markers ![alt](url)
        image_cache_dir: Directory to cache downloaded images
        font_size: Font size (default: 10.5pt)
        font_name: Font name (default: 宋体)
        bold: Whether text should be bold
        paragraph: Existing paragraph to add to (creates new if None)
        inline_images: If True, images are added inline in the same paragraph (for options)
                       If False, images are added in separate centered paragraphs
        paragraph_indent: If provided, apply this indent to image paragraphs
        
    Returns:
        List of paragraphs created (for tracking)
    """
    if font_size is None:
        font_size = Pt(10.5)
    
    created_paragraphs = []
    image_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
    image_matches = list(re.finditer(image_pattern, content))
    
    if not image_matches:
        # No images, just add text
        if paragraph is None:
            paragraph = doc.add_paragraph()
            created_paragraphs.append(paragraph)
        run = paragraph.add_run(content)
        run.font.size = font_size
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if bold:
            run.font.bold = True
        return created_paragraphs
    
    # Has images, process them
    last_pos = 0
    for img_match in image_matches:
        # Add text before image
        if img_match.start() > last_pos:
            text_part = content[last_pos:img_match.start()].strip()
            if text_part:
                if paragraph is None:
                    paragraph = doc.add_paragraph()
                    created_paragraphs.append(paragraph)
                run = paragraph.add_run(text_part)
                run.font.size = font_size
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                if bold:
                    run.font.bold = True
        
        # Process image
        img_url = img_match.group(2)
        img_alt = img_match.group(1) or ""
        
        # Download and insert image
        img_path = download_image(img_url, image_cache_dir)
        if img_path and img_path.exists():
            try:
                if inline_images and paragraph is not None:
                    # Add image inline in the same paragraph (for options)
                    run = paragraph.add_run()
                    run.add_picture(str(img_path), width=Inches(5))
                else:
                    # Add image in a new paragraph
                    img_para = doc.add_paragraph()
                    created_paragraphs.append(img_para)
                    if paragraph_indent is not None:
                        img_para.paragraph_format.left_indent = paragraph_indent
                    if not inline_images:
                        img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    img_para.paragraph_format.space_before = Pt(3)
                    img_para.paragraph_format.space_after = Pt(3)
                    img_run = img_para.add_run()
                    # Insert image with max width
                    img_run.add_picture(str(img_path), width=Inches(5 if inline_images else 6))
            except Exception as e:
                print(f"  Warning: Failed to insert image {img_url}: {str(e)}")
                # Add alt text as fallback
                if paragraph is None:
                    paragraph = doc.add_paragraph()
                    created_paragraphs.append(paragraph)
                run = paragraph.add_run(f"[图片: {img_alt}]" if img_alt else "[图片]")
                run.font.size = font_size
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        else:
            # Add alt text as fallback
            if paragraph is None:
                paragraph = doc.add_paragraph()
                created_paragraphs.append(paragraph)
            run = paragraph.add_run(f"[图片加载失败: {img_alt}]" if img_alt else "[图片加载失败]")
            run.font.size = font_size
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        
        last_pos = img_match.end()
    
    # Add remaining text after last image
    if last_pos < len(content):
        text_part = content[last_pos:].strip()
        if text_part:
            if paragraph is None:
                paragraph = doc.add_paragraph()
                created_paragraphs.append(paragraph)
            run = paragraph.add_run(text_part)
            run.font.size = font_size
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if bold:
                run.font.bold = True
    
    return created_paragraphs


def parse_markdown_to_docx(md_content: str, doc: Document, image_cache_dir: Optional[Path] = None):
    """Parse Markdown content and add it to a DOCX document with compact formatting.
    
    Args:
        md_content: Markdown content string
        doc: python-docx Document object
        image_cache_dir: Directory to cache downloaded images (optional)
    """
    lines = md_content.split('\n')
    i = 0
    
    # Image pattern: ![alt](url) or ![](url)
    image_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
    
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # Track if we're in an options section
    in_options_section = False
    option_index = 0
    
    while i < len(lines):
        line = lines[i].strip()
        original_line = lines[i]
        
        # Skip empty lines (we'll add minimal spacing manually)
        # Don't reset in_options_section on empty lines - options may have empty lines between them
        if not line:
            i += 1
            continue
        
        # H1 - Title
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(title)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            in_options_section = False
            option_index = 0
            i += 1
            continue
        
        # H2 - Section
        if line.startswith('## '):
            title = line[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            in_options_section = False
            option_index = 0
            i += 1
            continue
        
        # H3 - Question number
        if line.startswith('### '):
            title = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            in_options_section = False
            option_index = 0
            i += 1
            continue
        
        # Check for "**选项：**" to start options section
        if '**选项：**' in line:
            in_options_section = True
            option_index = 0
            # Still add the line
            match = re.match(r'\*\*(.+?)\*\*\s*(.*)', line)
            if match:
                bold_text = match.group(1)
                rest_text = match.group(2)
                p = doc.add_paragraph()
                run1 = p.add_run(bold_text)
                run1.font.bold = True
                run1.font.size = Pt(10.5)
                run1.font.name = '宋体'
                run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if rest_text:
                    run2 = p.add_run(rest_text)
                    run2.font.size = Pt(10.5)
                    run2.font.name = '宋体'
                    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            i += 1
            continue
        
        # Checkbox options - convert to ABCD format (only in options section)
        # Also check if this looks like an option (starts with - [ ] or - [x] after "**选项：**" section)
        if re.match(r'^- \[[ xX]\] ', original_line):
            # If we're in options section, or if this is the first option after "**选项：**"
            # (check previous lines for "**选项：**")
            if in_options_section or (i > 0 and any('**选项：**' in lines[j] for j in range(max(0, i-5), i))):
                if not in_options_section:
                    in_options_section = True
                    option_index = 0
            if in_options_section:
                # Extract option content
                content = re.sub(r'^- \[[ xX]\] ', '', original_line).strip()
                is_correct = '[x]' in original_line.lower() or '[X]' in original_line
                
                # Convert to ABCD (A=0, B=1, C=2, D=3, etc.)
                option_letter = chr(ord('A') + option_index)
                
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                
                run1 = p.add_run(f"{option_letter}. ")
                run1.font.size = Pt(10.5)
                run1.font.name = '宋体'
                run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                if is_correct:
                    run1.font.bold = True
                
                # Process content (may contain images)
                image_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
                image_matches = list(re.finditer(image_pattern, content))
                
                if not image_matches:
                    # No images, just add text
                    run2 = p.add_run(content)
                    run2.font.size = Pt(10.5)
                    run2.font.name = '宋体'
                    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    if is_correct:
                        run2.font.bold = True
                else:
                    # Has images, process them
                    last_pos = 0
                    for img_match in image_matches:
                        # Add text before image
                        if img_match.start() > last_pos:
                            text_part = content[last_pos:img_match.start()].strip()
                            if text_part:
                                run = p.add_run(text_part)
                                run.font.size = Pt(10.5)
                                run.font.name = '宋体'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                if is_correct:
                                    run.font.bold = True
                        
                        # Process image
                        img_url = img_match.group(2)
                        img_alt = img_match.group(1) or ""
                        
                        # Download and insert image
                        img_path = download_image(img_url, image_cache_dir)
                        if img_path and img_path.exists():
                            try:
                                # Add image inline in the same paragraph
                                run = p.add_run()
                                run.add_picture(str(img_path), width=Inches(4))
                            except Exception as e:
                                print(f"  Warning: Failed to insert image {img_url}: {str(e)}")
                                # Add alt text as fallback
                                run = p.add_run(f"[图片: {img_alt}]" if img_alt else "[图片]")
                                run.font.size = Pt(10.5)
                                run.font.name = '宋体'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        else:
                            # Add alt text as fallback
                            run = p.add_run(f"[图片加载失败: {img_alt}]" if img_alt else "[图片加载失败]")
                            run.font.size = Pt(10.5)
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        
                        last_pos = img_match.end()
                    
                    # Add remaining text after last image
                    if last_pos < len(content):
                        text_part = content[last_pos:].strip()
                        if text_part:
                            run = p.add_run(text_part)
                            run.font.size = Pt(10.5)
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            if is_correct:
                                run.font.bold = True
                    
                    # Add "(正确答案)" marker after images if this is the correct answer
                    if is_correct:
                        run = p.add_run(" （正确答案）")
                        run.font.size = Pt(10.5)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
                
                # Add "(正确答案)" marker for text-only options if this is the correct answer
                if is_correct and not image_matches:
                    run = p.add_run(" （正确答案）")
                    run.font.size = Pt(10.5)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
                
                option_index += 1
                i += 1
                continue
            else:
                # Not in options section, treat as regular list item
                content = re.sub(r'^- \[[ xX]\] ', '', original_line).strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(f"• {content}")
                run.font.size = Pt(10.5)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                i += 1
                continue
        
        # If we hit a non-option line after options, reset
        if in_options_section and not re.match(r'^- \[[ xX]\] ', original_line):
            in_options_section = False
            option_index = 0
        
        # Bold text (like **题目：**, **标准答案：**, etc.)
        if line.startswith('**'):
            # Extract bold text and rest
            match = re.match(r'\*\*(.+?)\*\*\s*(.*)', line)
            if match:
                bold_text = match.group(1)
                rest_text = match.group(2)
                p = doc.add_paragraph()
                run1 = p.add_run(bold_text)
                run1.font.bold = True
                run1.font.size = Pt(10.5)
                run1.font.name = '宋体'
                run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if rest_text:
                    # Process rest_text which may contain images
                    image_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
                    image_matches = list(re.finditer(image_pattern, rest_text))
                    
                    if not image_matches:
                        # No images, just add text
                        run2 = p.add_run(rest_text)
                        run2.font.size = Pt(10.5)
                        run2.font.name = '宋体'
                        run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    else:
                        # Has images, process them
                        last_pos = 0
                        for img_match in image_matches:
                            # Add text before image
                            if img_match.start() > last_pos:
                                text_part = rest_text[last_pos:img_match.start()].strip()
                                if text_part:
                                    run = p.add_run(text_part)
                                    run.font.size = Pt(10.5)
                                    run.font.name = '宋体'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            
                            # Process image
                            img_url = img_match.group(2)
                            img_alt = img_match.group(1) or ""
                            
                            # Download and insert image
                            img_path = download_image(img_url, image_cache_dir)
                            if img_path and img_path.exists():
                                try:
                                    # Add image inline in the same paragraph
                                    run = p.add_run()
                                    run.add_picture(str(img_path), width=Inches(5))
                                except Exception as e:
                                    print(f"  Warning: Failed to insert image {img_url}: {str(e)}")
                                    # Add alt text as fallback
                                    run = p.add_run(f"[图片: {img_alt}]" if img_alt else "[图片]")
                                    run.font.size = Pt(10.5)
                                    run.font.name = '宋体'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            else:
                                # Add alt text as fallback
                                run = p.add_run(f"[图片加载失败: {img_alt}]" if img_alt else "[图片加载失败]")
                                run.font.size = Pt(10.5)
                                run.font.name = '宋体'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            
                            last_pos = img_match.end()
                        
                        # Add remaining text after last image
                        if last_pos < len(rest_text):
                            text_part = rest_text[last_pos:].strip()
                            if text_part:
                                run = p.add_run(text_part)
                                run.font.size = Pt(10.5)
                                run.font.name = '宋体'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            i += 1
            continue
        
        # List items (starting with -)
        if line.startswith('- ') and not re.match(r'^- \[[ xX]\] ', original_line):
            content = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(f"• {content}")
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            i += 1
            continue
        
        # Horizontal rule
        if line.startswith('---'):
            # Add minimal spacing instead of a line
            in_options_section = False
            option_index = 0
            i += 1
            continue
        
        # Check for images in the line: ![alt](url)
        image_matches = list(re.finditer(image_pattern, original_line))
        if image_matches:
            # Split line by images and process text/image parts
            last_pos = 0
            
            for img_match in image_matches:
                # Add text before image
                if img_match.start() > last_pos:
                    text_part = original_line[last_pos:img_match.start()].strip()
                    if text_part:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        run = p.add_run(text_part)
                        run.font.size = Pt(10.5)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                # Process image
                img_url = img_match.group(2)
                img_alt = img_match.group(1) or ""
                
                # Download image
                img_path = download_image(img_url, image_cache_dir)
                if img_path and img_path.exists():
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(6)
                        run = p.add_run()
                        # Insert image with max width of 6 inches
                        run.add_picture(str(img_path), width=Inches(6))
                    except Exception as e:
                        print(f"  Warning: Failed to insert image {img_url}: {str(e)}")
                        # Add alt text as fallback
                        if img_alt:
                            p = doc.add_paragraph()
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                            run = p.add_run(f"[图片: {img_alt}]")
                            run.font.size = Pt(10.5)
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                else:
                    # Add alt text as fallback
                    if img_alt:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        run = p.add_run(f"[图片加载失败: {img_alt}]")
                        run.font.size = Pt(10.5)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                last_pos = img_match.end()
            
            # Add remaining text after last image
            if last_pos < len(original_line):
                text_part = original_line[last_pos:].strip()
                if text_part:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(text_part)
                    run.font.size = Pt(10.5)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # If line was only images, we've already processed it
            i += 1
            continue
        
        # Regular text (process text that may contain images inline, but already handled above)
        # Check if line contains image markers but wasn't handled (shouldn't happen, but safety check)
        if re.search(image_pattern, line):
            # Remove image markers and process as text
            line = re.sub(image_pattern, '', line).strip()
            if not line:
                i += 1
                continue
        
        # Regular text
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        i += 1


def convert_markdown_to_docx(md_file: Path, docx_file: Path, image_cache_dir: Optional[Path] = None):
    """Convert a single Markdown file to DOCX.
    
    Args:
        md_file: Path to Markdown file
        docx_file: Path to output DOCX file
        image_cache_dir: Directory to cache downloaded images (optional)
    """
    # Read Markdown content
    with open(md_file, "r", encoding="utf-8") as f:
        md_content = f.read()
    
    # Create new document
    doc = Document()
    
    # Set page margins (compact)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # Parse and add content (with image support)
    parse_markdown_to_docx(md_content, doc, image_cache_dir=image_cache_dir)
    
    # Save document
    doc.save(docx_file)


def main():
    """Main function to convert Markdown files to DOCX."""
    parser = argparse.ArgumentParser(
        description="Convert Markdown files to DOCX format"
    )
    parser.add_argument(
        "--input",
        type=str,
        default="output/markdown",
        help="Input directory containing Markdown files (default: output/markdown)",
    )
    parser.add_argument(
        "--output",
        type=str,
        default="output/docx",
        help="Output directory for DOCX files (default: output/docx)",
    )
    parser.add_argument(
        "--courses",
        nargs="+",
        default=None,
        help="List of course names to process (only process these courses)",
    )
    parser.add_argument(
        "--courses-file",
        type=str,
        default=None,
        help="JSON file containing list of course names to process",
    )
    parser.add_argument(
        "--image-cache",
        type=str,
        default=None,
        help="Directory to cache downloaded images (default: output/images)",
    )
    
    args = parser.parse_args()
    
    # Load courses from file if specified
    selected_courses = args.courses
    if args.courses_file:
        import json
        courses_file = Path(args.courses_file)
        if courses_file.exists():
            try:
                with open(courses_file, "r", encoding="utf-8") as f:
                    selected_courses = json.load(f)
            except Exception as e:
                print(f"Failed to load courses file: {str(e)}")
                selected_courses = None
    
    markdown_dir = Path(args.input)
    docx_dir = Path(args.output)
    
    # Set up image cache directory
    if args.image_cache:
        image_cache_dir = Path(args.image_cache)
    else:
        image_cache_dir = docx_dir.parent / "images"
    image_cache_dir.mkdir(parents=True, exist_ok=True)
    
    if not markdown_dir.exists():
        print(f"Input directory not found: {markdown_dir}")
        print("  Please run convert_markdown.py first to generate Markdown files")
        return
    
    try:
        # Find all Markdown files
        print("Scanning for Markdown files...")
        courses_data = find_markdown_files(markdown_dir)
        
        if not courses_data:
            print("No Markdown files found")
            return
        
        total_courses = len(courses_data)
        total_files = sum(len(files) for files in courses_data.values())
        print(f"Found {total_courses} courses with {total_files} Markdown files")
        
        # Filter by selected courses if specified
        if selected_courses:
            filtered_courses_data = {
                name: files for name, files in courses_data.items()
                if name in selected_courses
            }
            if filtered_courses_data:
                print(f"Processing {len(filtered_courses_data)} selected course(s): {', '.join(filtered_courses_data.keys())}")
                courses_data = filtered_courses_data
            else:
                print("No matching courses found in selected list")
                return
        
        print()
        
        # Convert to DOCX
        print("=" * 60)
        print("Converting to DOCX...")
        print("=" * 60)
        
        total_exported = 0
        total_skipped = 0
        total_errors = 0
        
        for course_name, markdown_files in courses_data.items():
            # Create course directory in DOCX output
            safe_course_name = course_name.replace("/", "_").replace("\\", "_")
            course_docx_dir = docx_dir / safe_course_name
            course_docx_dir.mkdir(parents=True, exist_ok=True)
            
            course_exported = 0
            course_skipped = 0
            
            for md_file in markdown_files:
                try:
                    # Generate DOCX filename
                    docx_filename = md_file.stem + ".docx"
                    docx_file = course_docx_dir / docx_filename
                    
                    # Skip if DOCX file already exists
                    if docx_file.exists():
                        course_skipped += 1
                        total_skipped += 1
                        continue
                    
                    # Convert to DOCX (with image support)
                    convert_markdown_to_docx(md_file, docx_file, image_cache_dir=image_cache_dir)
                    course_exported += 1
                    total_exported += 1
                except Exception as e:
                    total_errors += 1
                    print(f"  Failed to convert {md_file.name}: {str(e)}")
            
            status = f"{course_exported} exported"
            if course_skipped > 0:
                status += f", {course_skipped} skipped"
            print(f"  {course_name}: {status}")
        
        print()
        print("=" * 60)
        print(f"   Total DOCX files exported: {total_exported}")
        if total_skipped > 0:
            print(f"   Total DOCX files skipped: {total_skipped} (already exist)")
        if total_errors > 0:
            print(f"   Errors: {total_errors}")
        print(f"   DOCX files saved to: {docx_dir}")
        print("=" * 60)
    
    except KeyboardInterrupt:
        print("\n\nInterrupted by user")
    except Exception as e:
        print(f"\nError: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
